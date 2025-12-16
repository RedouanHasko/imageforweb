# app.py
# Professional lightweight web app to convert PNG/JPEG images to optimized web-ready images
# FIXED VERSION: Uses Flask instead of FastAPI to avoid SSL module dependency issues
# Stack: Flask + Pillow + clean modern UI (Tailwind CDN)

from flask import Flask, request, send_file, Response, jsonify
from PIL import Image
try:
  from pdf2image import convert_from_bytes
except Exception:
  convert_from_bytes = None
try:
  import fitz  # PyMuPDF
except Exception:
  fitz = None
import io
import zipfile
import threading
import uuid
import tempfile
import os
import time
import subprocess
import shutil
try:
  from pdf2docx import Converter
except Exception:
  Converter = None
try:
  import pytesseract
except Exception:
  pytesseract = None
try:
  from docx import Document
except Exception:
  Document = None

# Detect LibreOffice `soffice` executable if available for high-fidelity PDF->DOCX
SOFFICE = shutil.which('soffice') or shutil.which('soffice.exe')

# In-memory job store. For production consider a persistent queue/storage.
jobs = {}
jobs_lock = threading.Lock()

app = Flask(__name__)

@app.route("/", methods=["GET"])
def index():
    return """
<!DOCTYPE html>
<html lang='en'>
<head>
  <meta charset='UTF-8'>
  <meta name='viewport' content='width=device-width, initial-scale=1.0'>
  <title>Image Optimizer</title>
  <script src='https://cdn.tailwindcss.com'></script>
  <style>
    .drag-active { border: 2px dashed #6366f1 !important; background: rgba(99,102,241,0.06) !important; }
  </style>
</head>
<body class='bg-gradient-to-b from-zinc-900 via-zinc-900 to-zinc-800 text-white min-h-screen flex items-center justify-center p-6'>
  <div class='w-full max-w-5xl grid grid-cols-12 gap-8'>
    <div class='col-span-7 bg-zinc-800/70 backdrop-blur rounded-3xl p-8 shadow-xl'>
      <div class='flex items-center justify-between mb-6'>
        <div>
          <h1 class='text-4xl font-extrabold'>Image Optimizer</h1>
          <p class='text-zinc-400 mt-1'>Fast conversions — images, PDFs, combined PDFs.</p>
        </div>
        <div class='text-right text-sm text-zinc-400'>
          <div>Drag & drop or click to select</div>
          <div class='mt-2'>Formats: WebP, JPEG, PNG, AVIF, PDF</div>
        </div>
      </div>

      <form id='uploadForm' action='/optimize' method='post' enctype='multipart/form-data'>
        <div id='dropArea' class='relative mb-6 flex flex-col items-center justify-center border-2 border-dashed border-zinc-600 rounded-3xl p-8 min-h-[14rem] transition'>
          <svg class='w-14 h-14 mb-3 text-indigo-400' fill='none' stroke='currentColor' stroke-width='1.5' viewBox='0 0 24 24'><path stroke-linecap='round' stroke-linejoin='round' d='M7 16V4a1 1 0 0 1 1-1h8a1 1 0 0 1 1 1v12m-5 4v-4m0 0l-2 2m2-2l2 2'/></svg>
          <div class='text-center'>
            <div class='text-lg font-medium text-zinc-200'>Drop files here</div>
            <div class='text-sm text-zinc-500'>Images (PNG/JPEG) or PDFs — click to open file picker</div>
          </div>
          <input id='fileInput' type='file' name='files' multiple accept='image/png,image/jpeg,application/pdf' class='absolute inset-0 opacity-0 cursor-pointer' />
        </div>

        <div class='grid grid-cols-2 gap-4 mb-4'>
          <div>
            <label class='text-zinc-300 text-sm'>Format</label>
            <select name='format' class='mt-2 w-full bg-zinc-700 rounded-lg p-3'>
              <option value='webp'>WebP (Recommended)</option>
              <option value='jpeg'>JPEG</option>
              <option value='png'>PNG</option>
              <option value='avif'>AVIF (optional)</option>
              <option value='pdf'>PDF (images → single-page PDF or PDF passthrough)</option>
              <option value='docx'>DOCX (editable Word document)</option>
            </select>
            <div id='pdfHint' class='text-xs text-zinc-500 mt-2 hidden'>Higher <strong>Quality</strong> increases rendering DPI and file size for PDFs.</div>
          </div>

          <div>
            <label class='text-zinc-300 text-sm'>Quality</label>
            <div class='mt-2 flex items-center gap-3'>
              <input id='qualityRange' name='quality' type='range' min='10' max='100' value='85' class='flex-1' />
              <div id='qualityVal' class='w-14 text-right text-zinc-300'>85</div>
            </div>
          </div>

          <div>
            <label class='text-zinc-300 text-sm'>Max Width (px)</label>
            <input name='max_width' type='number' min='0' placeholder='1920' class='mt-2 w-full bg-zinc-700 rounded-lg p-3' />
          </div>

          <div>
            <label class='text-zinc-300 text-sm'>Max Height (px)</label>
            <input name='max_height' type='number' min='0' placeholder='1080' class='mt-2 w-full bg-zinc-700 rounded-lg p-3' />
          </div>
        </div>

        <div class='flex items-center gap-4 mb-6'>
          <label id='combineWrap' class='hidden items-center gap-3 text-zinc-300'>
            <input id='combinePdf' name='combine_pdf' type='checkbox' value='1' class='h-5 w-5 accent-indigo-500 rounded' />
            <span class='select-none'>Combine into a single PDF</span>
          </label>
          <label id='ocrWrap' class='items-center gap-3 text-zinc-300'>
            <input id='ocrCheckbox' name='ocr' type='checkbox' value='1' class='h-5 w-5 accent-indigo-500 rounded' />
            <span class='select-none'>OCR (make scanned PDFs/images editable)</span>
          </label>
          <label id='preserveWrap' class='hidden items-center gap-3 text-zinc-300'>
            <input id='preserveLayout' name='preserve_layout' type='checkbox' value='1' class='h-5 w-5 accent-indigo-500 rounded' />
            <span class='select-none'>Preserve layout (LibreOffice)</span>
          </label>
          <div class='flex-1'></div>
        </div>

        <div class='flex gap-3'>
          <button id='submitBtn' type='submit' class='flex-1 bg-indigo-600 hover:bg-indigo-700 rounded-2xl py-3 text-lg font-semibold'>Start</button>
          <button id='clearBtn' type='button' class='w-40 bg-zinc-600 hover:bg-zinc-700 rounded-2xl py-3 text-sm'>Clear</button>
        </div>
      </form>
    </div>

    <div class='col-span-5'>
      <div class='bg-zinc-800/60 rounded-3xl p-6 shadow-lg h-full flex flex-col'>
        <div class='mb-4 flex items-center justify-between'>
          <div class='text-sm text-zinc-400'>Selected files</div>
          <div id='imageCount' class='text-sm text-zinc-400'>No files selected</div>
        </div>
        <div id='fileList' class='flex-1 overflow-auto space-y-3 mb-4'></div>

        <div id='progressArea' class='hidden'>
          <div class='w-full bg-zinc-700 rounded-xl h-4 mb-3 overflow-hidden'>
            <div id='progressBar' class='bg-indigo-500 h-4 rounded-xl transition-all' style='width:0%'></div>
          </div>
          <div class='flex justify-between text-sm'>
            <div id='progressText'>0%</div>
            <div id='processedCount' class='text-zinc-400'></div>
          </div>
          <div id='statusMsg' class='mt-2 text-sm'></div>
        </div>
      </div>
    </div>
  </div>

  <script>
    // Drag and drop logic
    const dropArea = document.getElementById('dropArea');
    const fileInput = document.getElementById('fileInput');
    const uploadForm = document.getElementById('uploadForm');
    const imageCount = document.getElementById('imageCount');
    const processedCount = document.getElementById('processedCount');
    const progressArea = document.getElementById('progressArea');
    const progressBar = document.getElementById('progressBar');
    const progressText = document.getElementById('progressText');
    const statusMsg = document.getElementById('statusMsg');
    let files = [];

    function updateImageCount() {
      if (files.length === 0) {
        imageCount.textContent = 'No files selected';
      } else {
        imageCount.textContent = files.length + ' file' + (files.length > 1 ? 's' : '') + ' selected';
      }
    }

    function renderFileList(serverItems) {
      const list = document.getElementById('fileList');
      if (!files || files.length === 0) { list.innerHTML = ''; return; }
      let html = "";
      files.forEach((f, i) => {
        const sizeKB = Math.round(f.size / 1024);
        let status = 'queued';
        let out = '';
        if (serverItems && serverItems[i]) {
          status = serverItems[i].status || status;
          out = serverItems[i].out_name ? (' → ' + serverItems[i].out_name) : '';
        }
        const color = status === 'done' ? 'text-green-400' : status === 'processing' ? 'text-indigo-300' : status === 'error' ? 'text-red-400' : 'text-zinc-400';
        html += `<div class='flex items-center justify-between bg-zinc-700 p-3 rounded-2xl'>`;
        html += `<div class='truncate'><strong class='text-sm'>${f.name}</strong> <span class='text-zinc-400 text-sm'>(${sizeKB} KB)</span>${out}</div>`;
        html += `<div class='ml-4 ${color} text-sm'>${status}</div>`;
        html += `</div>`;
      });
      list.innerHTML = html;
    }

    // Quality slider sync
    const qualityRange = document.getElementById('qualityRange');
    const qualityVal = document.getElementById('qualityVal');
    if (qualityRange && qualityVal) { qualityRange.addEventListener('input', (e) => { qualityVal.textContent = e.target.value; }); }

    // Show / hide Combine / Preserve options depending on selected format
    const formatSelect = uploadForm.querySelector('select[name="format"]');
    const combineWrap = document.getElementById('combineWrap');
    const preserveWrap = document.getElementById('preserveWrap');
    if (formatSelect && combineWrap) {
      function updateCombineVisibility() {
        const pdfHint = document.getElementById('pdfHint');
        if (formatSelect.value === 'pdf') {
          combineWrap.classList.remove('hidden');
          if (pdfHint) pdfHint.classList.remove('hidden');
        } else {
          combineWrap.classList.add('hidden');
          if (pdfHint) pdfHint.classList.add('hidden');
          const cb = document.getElementById('combinePdf'); if (cb) cb.checked = false;
        }
        // Show preserve layout option only for DOCX output
        if (preserveWrap) {
          if (formatSelect.value === 'docx') {
            preserveWrap.classList.remove('hidden');
          } else {
            preserveWrap.classList.add('hidden');
            const pb = document.getElementById('preserveLayout'); if (pb) pb.checked = false;
          }
        }
      }
      formatSelect.addEventListener('change', updateCombineVisibility); updateCombineVisibility();
    }

    // Clear / reset UI and selected files
    const clearBtn = document.getElementById('clearBtn');
    if (clearBtn) { clearBtn.addEventListener('click', () => { files = []; const dt = new DataTransfer(); fileInput.files = dt.files; try { fileInput.value = ''; } catch (e) {} updateImageCount(); renderFileList(); progressArea.classList.add('hidden'); progressBar.style.width = '0%'; progressText.textContent = '0%'; statusMsg.textContent = ''; processedCount.textContent = ''; const fmt = uploadForm.querySelector('select[name="format"]'); if (fmt) fmt.value = 'webp'; if (qualityRange) { qualityRange.value = 85; qualityVal.textContent = '85'; } const mw = uploadForm.querySelector('input[name="max_width"]'); const mh = uploadForm.querySelector('input[name="max_height"]'); if (mw) mw.value = ''; if (mh) mh.value = ''; }); }

    dropArea.addEventListener('click', (e) => { if (e.target === fileInput) return; fileInput.click(); });
    dropArea.addEventListener('dragover', e => { e.preventDefault(); dropArea.classList.add('drag-active'); });
    dropArea.addEventListener('dragleave', e => { e.preventDefault(); dropArea.classList.remove('drag-active'); });
    dropArea.addEventListener('drop', e => { e.preventDefault(); dropArea.classList.remove('drag-active'); const dropped = Array.from(e.dataTransfer.files).filter(f => f.type.startsWith('image/') || f.type === 'application/pdf' || (f.name && f.name.toLowerCase().endsWith('.pdf'))); dropped.forEach(df => { if (!files.some(f => f.name === df.name && f.size === df.size)) files.push(df); }); const dt = new DataTransfer(); files.forEach(f => dt.items.add(f)); fileInput.files = dt.files; updateImageCount(); renderFileList(); });
    fileInput.addEventListener('change', e => { const newly = Array.from(fileInput.files).filter(f => f.type.startsWith('image/') || f.type === 'application/pdf' || (f.name && f.name.toLowerCase().endsWith('.pdf'))); newly.forEach(nf => { if (!files.some(f => f.name === nf.name && f.size === nf.size)) files.push(nf); }); const dt2 = new DataTransfer(); files.forEach(f => dt2.items.add(f)); fileInput.files = dt2.files; updateImageCount(); renderFileList(); });

    // Progress and status logic
    uploadForm.addEventListener('submit', function(e) { e.preventDefault(); statusMsg.textContent = ''; statusMsg.className = ''; if (files.length === 0) { statusMsg.textContent = 'Please select at least one file.'; statusMsg.className = 'text-red-400'; return; } progressArea.classList.remove('hidden'); progressBar.style.width = '0%'; progressText.textContent = '0%'; processedCount.textContent = '';
      const formData = new FormData(); const format = uploadForm.querySelector('select[name="format"]').value; const maxw = uploadForm.querySelector('input[name="max_width"]').value || ''; const maxh = uploadForm.querySelector('input[name="max_height"]').value || ''; const qualityControl = document.getElementById('qualityRange'); const qualityValToSend = qualityControl ? qualityControl.value : (uploadForm.querySelector('input[name="quality"]') ? uploadForm.querySelector('input[name="quality"]').value : '85'); formData.append('format', format); formData.append('quality', qualityValToSend); formData.append('max_width', maxw); formData.append('max_height', maxh); files.forEach(f => formData.append('files', f)); const combineCheckbox = document.getElementById('combinePdf'); formData.append('combine_pdf', combineCheckbox && combineCheckbox.checked ? '1' : '0'); const ocrCheckbox = document.getElementById('ocrCheckbox'); formData.append('ocr', ocrCheckbox && ocrCheckbox.checked ? '1' : '0'); const preserveCheckbox = document.getElementById('preserveLayout'); formData.append('preserve_layout', preserveCheckbox && preserveCheckbox.checked ? '1' : '0');
      const xhr = new XMLHttpRequest(); xhr.open('POST', '/start', true); xhr.responseType = 'json'; xhr.upload.onprogress = function(e) { if (e.lengthComputable) { const percent = Math.round((e.loaded / e.total) * 100 * 0.4); progressBar.style.width = percent + '%'; progressText.textContent = percent + '%'; } };
      xhr.onload = function() { if (xhr.status === 200 && xhr.response && xhr.response.job_id) { const jobId = xhr.response.job_id; const total = xhr.response.total || files.length; imageCount.textContent = total + ' file' + (total > 1 ? 's' : '') + ' queued'; const poll = setInterval(async () => { try { const res = await fetch(`/status/${jobId}`); const data = await res.json(); if (data.error) { clearInterval(poll); statusMsg.textContent = 'Error: ' + data.error; statusMsg.className = 'text-red-400'; return; } if (data.items) { renderFileList(data.items); } const proc = data.processed || 0; const tot = data.total || total; const serverPercent = Math.round((proc / (tot || 1)) * 100 * 0.6); const uploadPercent = parseInt(progressText.textContent) || 0; const combined = Math.min(100, uploadPercent + serverPercent); progressBar.style.width = combined + '%'; progressText.textContent = combined + '%'; processedCount.textContent = proc + ' / ' + tot + ' processed'; if (data.status === 'done') { clearInterval(poll); progressBar.style.width = '100%'; progressText.textContent = '100%'; statusMsg.textContent = 'Done!'; statusMsg.className = 'text-green-400'; fetch(`/download/${jobId}`).then(r => r.blob()).then(blob => { const url = window.URL.createObjectURL(blob); const a = document.createElement('a'); a.href = url; a.download = 'optimized_images.zip'; document.body.appendChild(a); a.click(); setTimeout(() => { window.URL.revokeObjectURL(url); document.body.removeChild(a); }, 100); }).catch(err => { statusMsg.textContent = 'Download error'; statusMsg.className = 'text-red-400'; }); } else if (data.status === 'error') { clearInterval(poll); statusMsg.textContent = 'Error: ' + data.error; statusMsg.className = 'text-red-400'; } } catch (err) { clearInterval(poll); statusMsg.textContent = 'Status error'; statusMsg.className = 'text-red-400'; } }, 600); } else { statusMsg.textContent = 'Failed to start job.'; statusMsg.className = 'text-red-400'; } };
      xhr.onerror = function() { statusMsg.textContent = 'Network error.'; statusMsg.className = 'text-red-400'; };
      xhr.send(formData);
    });
  </script>
</body>
</html>
"""

@app.route("/optimize", methods=["POST"])
def optimize_images():
    files = request.files.getlist("files")
    output_format = request.form.get("format", "webp")

    if not files:
        return Response("No files uploaded", status=400)

    zip_buffer = io.BytesIO()
    processed = 0
    try:
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file in files:
                try:
                    img = Image.open(file.stream).convert('RGB')
                    output = io.BytesIO()
                    if output_format == 'webp':
                        img.save(output, format='WEBP', quality=85, method=6)
                        ext = 'webp'
                    else:
                        img.save(output, format='JPEG', quality=85, optimize=True, progressive=True)
                        ext = 'jpg'
                    filename = file.filename.rsplit('.', 1)[0]
                    zipf.writestr(f"{filename}.{ext}", output.getvalue())
                    processed += 1
                except Exception as e:
                    # If one image fails, skip and continue
                    continue
        zip_buffer.seek(0)
        if processed == 0:
            return Response("All images failed to process.", status=500)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='optimized_images.zip'
        )
    except Exception as e:
        return Response(f"Error: {str(e)}", status=500)

# =============================
# BASIC TEST CASES (MANUAL)
# =============================
# 1. Upload multiple PNG files → receive ZIP with .webp images
# 2. Upload multiple JPEG files → receive ZIP with optimized .jpg images
# 3. Upload mixed PNG/JPEG → all converted correctly
# 4. Upload nothing → returns HTTP 400

# Run with:
# pip install flask pillow


def _process_job(job_id, file_blobs, output_format, quality, max_w, max_h, combine_pdf=False, ocr=False, preserve_layout=False):
  """Background worker to process images for a job and write a ZIP to disk.
  Updates the shared `jobs` dict with progress.
  """
  try:
    zip_buffer = io.BytesIO()
    processed = 0
    total = len(file_blobs)
    combined_pages = [] if (output_format == 'pdf' and combine_pdf) else None
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
      for idx, (name, blob) in enumerate(file_blobs):
        # mark processing
        with jobs_lock:
          if job_id in jobs and 'items' in jobs[job_id] and idx < len(jobs[job_id]['items']):
            jobs[job_id]['items'][idx]['status'] = 'processing'
        try:
          # Handle PDF inputs before attempting to open as images
          in_name_l = name.lower()
          if in_name_l.endswith('.pdf'):
            if not convert_from_bytes and not fitz:
              raise RuntimeError('pdf2image not available; install pdf2image and poppler, or install PyMuPDF as a fallback')
            # Choose rendering DPI based on requested quality to preserve visual detail
            try:
              qval = int(quality)
            except Exception:
              qval = 85
            if qval >= 95:
              dpi = 400
            elif qval >= 90:
              dpi = 300
            elif qval >= 75:
              dpi = 200
            else:
              dpi = 150
            try:
              # prefer pdf2image when available for accurate rendering
              if convert_from_bytes:
                pages = convert_from_bytes(blob, dpi=dpi)
              else:
                raise RuntimeError('pdf2image not available')
            except Exception as e:
              # pdf2image failed (often due to missing poppler). Try PyMuPDF (fitz) as a fallback if available.
              if fitz:
                try:
                  doc = fitz.open(stream=blob, filetype='pdf')
                  pages = []
                  # Use matrix zoom to match DPI used by pdf2image
                  zoom = dpi / 72.0
                  mat = fitz.Matrix(zoom, zoom)
                  for p in doc:
                    pix = p.get_pixmap(matrix=mat, alpha=False)
                    img = Image.frombytes('RGB', (pix.width, pix.height), pix.samples)
                    pages.append(img)
                except Exception as e2:
                  msg = str(e2)
                  with jobs_lock:
                    if job_id in jobs and 'items' in jobs[job_id] and idx < len(jobs[job_id]['items']):
                      jobs[job_id]['items'][idx]['status'] = 'error'
                      jobs[job_id]['items'][idx]['error'] = msg + ' — Try installing poppler or PyMuPDF.'
                      jobs[job_id]['processed'] = processed
                  continue
              else:
                msg = str(e)
                if 'poppler' in msg.lower() or 'page count' in msg.lower():
                  msg += ' — On Windows install Poppler (add its `bin` to PATH). See https://github.com/oschwartz10612/poppler-windows/releases'
                msg += ' — Or install PyMuPDF (`pip install pymupdf`) to enable a fallback renderer.'
                with jobs_lock:
                  if job_id in jobs and 'items' in jobs[job_id] and idx < len(jobs[job_id]['items']):
                    jobs[job_id]['items'][idx]['status'] = 'error'
                    jobs[job_id]['items'][idx]['error'] = msg
                    jobs[job_id]['processed'] = processed
                continue
            # write each page as separate file (or collect pages for combined PDF / DOCX)
            if output_format in ('webp', 'jpeg', 'png', 'avif'):
              for pidx, page in enumerate(pages, start=1):
                pout = io.BytesIO()
                fmt = output_format
                if fmt == 'jpeg': save_fmt = 'JPEG'
                elif fmt == 'png': save_fmt = 'PNG'
                else: save_fmt = 'WEBP'
                page = page.convert('RGB')
                if save_fmt == 'WEBP':
                  page.save(pout, format='WEBP', quality=quality, method=6)
                  ext = 'webp'
                elif save_fmt == 'PNG':
                  page.save(pout, format='PNG', optimize=True)
                  ext = 'png'
                else:
                  page.save(pout, format='JPEG', quality=quality, optimize=True, progressive=True)
                  ext = 'jpg'
                if combined_pages is not None:
                  combined_pages.append(page)
                else:
                  zipf.writestr(f"{os.path.splitext(name)[0]}_page{pidx}.{ext}", pout.getvalue())
              processed += 1
              with jobs_lock:
                if job_id in jobs and 'items' in jobs[job_id] and idx < len(jobs[job_id]['items']):
                  jobs[job_id]['items'][idx]['status'] = 'done'
                  jobs[job_id]['items'][idx]['out_name'] = f"{os.path.splitext(name)[0]}_pages.{output_format}"
              continue
            elif output_format == 'docx':
              # Try high-fidelity LibreOffice conversion when requested
              filename = os.path.splitext(name)[0]
              if preserve_layout and SOFFICE:
                try:
                  tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                  tmp_pdf.write(blob)
                  tmp_pdf.close()
                  outdir = tempfile.mkdtemp()
                  cmd = [SOFFICE, '--headless', '--convert-to', 'docx', '--outdir', outdir, tmp_pdf.name]
                  subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                  out_path = os.path.join(outdir, os.path.splitext(os.path.basename(tmp_pdf.name))[0] + '.docx')
                  with open(out_path, 'rb') as rf:
                    data = rf.read()
                  zipf.writestr(f"{filename}.docx", data)
                  # cleanup
                  try:
                    os.unlink(tmp_pdf.name)
                  except Exception:
                    pass
                  try:
                    os.remove(out_path)
                  except Exception:
                    pass
                  try:
                    os.rmdir(outdir)
                  except Exception:
                    pass
                except Exception as e:
                  # If soffice fails, fall back to the next available method
                  try:
                    if os.path.exists(tmp_pdf.name):
                      os.unlink(tmp_pdf.name)
                  except Exception:
                    pass
                else:
                  processed += 1
                  with jobs_lock:
                    if job_id in jobs and 'items' in jobs[job_id] and idx < len(jobs[job_id]['items']):
                      jobs[job_id]['items'][idx]['status'] = 'done'
                      jobs[job_id]['items'][idx]['out_name'] = f"{os.path.splitext(name)[0]}.docx"
                  continue

              # Try using pdf2docx for digital PDFs when OCR is not requested
              if Converter and not ocr:
                try:
                  tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                  tmp_pdf.write(blob)
                  tmp_pdf.close()
                  tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                  tmp_out.close()
                  conv = Converter(tmp_pdf.name)
                  conv.convert(tmp_out.name)
                  conv.close()
                  with open(tmp_out.name, 'rb') as rf:
                    data = rf.read()
                  zipf.writestr(f"{filename}.docx", data)
                except Exception as e:
                  with jobs_lock:
                    jobs[job_id]['items'][idx]['status'] = 'error'
                    jobs[job_id]['items'][idx]['error'] = str(e)
                finally:
                  try:
                    os.unlink(tmp_pdf.name)
                  except Exception:
                    pass
                  try:
                    os.unlink(tmp_out.name)
                  except Exception:
                    pass
              else:
                # Prefer PyMuPDF text extraction for digital PDFs (no OCR), otherwise fall back to OCR on images
                if not Document:
                  with jobs_lock:
                    jobs[job_id]['items'][idx]['status'] = 'error'
                    jobs[job_id]['items'][idx]['error'] = 'python-docx not installed'
                  continue
                if not ocr and fitz:
                  try:
                    doc = Document()
                    pdfdoc = fitz.open(stream=blob, filetype='pdf')
                    for p in pdfdoc:
                      try:
                        text = p.get_text('text')
                      except Exception:
                        text = ''
                      if text:
                        for line in text.splitlines():
                          if line.strip():
                            doc.add_paragraph(line)
                      # Optionally embed a low-res image of the page for reference
                      try:
                        pix = p.get_pixmap(alpha=False)
                        tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                        pix.save(tmp_img.name)
                        tmp_img.close()
                        doc.add_picture(tmp_img.name)
                        os.unlink(tmp_img.name)
                      except Exception:
                        pass
                    outbuf = io.BytesIO()
                    doc.save(outbuf)
                    zipf.writestr(f"{filename}.docx", outbuf.getvalue())
                  except Exception as e:
                    with jobs_lock:
                      jobs[job_id]['items'][idx]['status'] = 'error'
                      jobs[job_id]['items'][idx]['error'] = str(e)
                    continue
                else:
                  # OCR-based fallback: render pages and extract text with pytesseract
                  doc = Document()
                  for pidx, page in enumerate(pages, start=1):
                    text = ''
                    if pytesseract:
                      try:
                        text = pytesseract.image_to_string(page)
                      except Exception:
                        text = ''
                    if text:
                      for line in text.splitlines():
                        if line.strip():
                          doc.add_paragraph(line)
                    # embed page image as well
                    try:
                      tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                      page.save(tmp_img.name, format='PNG')
                      tmp_img.close()
                      doc.add_picture(tmp_img.name)
                      os.unlink(tmp_img.name)
                    except Exception:
                      pass
                  outbuf = io.BytesIO()
                  try:
                    doc.save(outbuf)
                    zipf.writestr(f"{filename}.docx", outbuf.getvalue())
                  except Exception as e:
                    with jobs_lock:
                      jobs[job_id]['items'][idx]['status'] = 'error'
                      jobs[job_id]['items'][idx]['error'] = str(e)
                    continue
              processed += 1
              with jobs_lock:
                if job_id in jobs and 'items' in jobs[job_id] and idx < len(jobs[job_id]['items']):
                  jobs[job_id]['items'][idx]['status'] = 'done'
                  jobs[job_id]['items'][idx]['out_name'] = f"{os.path.splitext(name)[0]}.docx"
              continue

          img = Image.open(io.BytesIO(blob)).convert('RGB')
          # Resize if needed while preserving aspect ratio
          if max_w or max_h:
            mw = int(max_w) if max_w else None
            mh = int(max_h) if max_h else None
            if mw or mh:
              target = (mw or img.width, mh or img.height)
              img.thumbnail(target, Image.LANCZOS)
          out = io.BytesIO()

          if output_format == 'webp':
            img.save(out, format='WEBP', quality=quality, method=6)
            ext = 'webp'
          elif output_format == 'png':
            img.save(out, format='PNG', optimize=True)
            ext = 'png'
          elif output_format == 'avif':
            try:
              img.save(out, format='AVIF', quality=quality)
              ext = 'avif'
            except Exception:
              img.save(out, format='WEBP', quality=quality, method=6)
              ext = 'webp'
          elif output_format == 'pdf':
            # images -> single-page PDF per image
            img.save(out, format='PDF', resolution=72)
            ext = 'pdf'
          else:
            img.save(out, format='JPEG', quality=quality, optimize=True, progressive=True)
            ext = 'jpg'

          filename = os.path.splitext(name)[0]
          if combined_pages is not None and output_format == 'pdf':
            # collect the PIL image for later combined PDF
            combined_pages.append(img)
          else:
            zipf.writestr(f"{filename}.{ext}", out.getvalue())
          processed += 1
          # update per-item status
          with jobs_lock:
            if job_id in jobs and 'items' in jobs[job_id] and idx < len(jobs[job_id]['items']):
              jobs[job_id]['items'][idx]['status'] = 'done'
              jobs[job_id]['items'][idx]['out_name'] = f"{filename}.{ext}"
            jobs[job_id]['processed'] = processed
        except Exception as e:
          with jobs_lock:
            if job_id in jobs and 'items' in jobs[job_id] and idx < len(jobs[job_id]['items']):
              jobs[job_id]['items'][idx]['status'] = 'error'
              jobs[job_id]['items'][idx]['error'] = str(e)
            jobs[job_id]['processed'] = processed
          continue

    # If combining into a single PDF, create combined PDF and add to ZIP
    if combined_pages is not None and len(combined_pages) > 0:
      try:
        first, rest = combined_pages[0], combined_pages[1:]
        outbuf = io.BytesIO()
        first.save(outbuf, format='PDF', save_all=True, append_images=rest)
        # The original `zipf` was closed when exiting the `with` above,
        # so reopen the in-memory ZIP in append mode and write the combined PDF.
        zip_buffer.seek(0)
        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zipf2:
          zipf2.writestr('combined.pdf', outbuf.getvalue())
        # mark items' out_name to combined.pdf
        with jobs_lock:
          if job_id in jobs and 'items' in jobs[job_id]:
            for it in jobs[job_id]['items']:
              it['out_name'] = 'combined.pdf'
      except Exception as e:
        with jobs_lock:
          jobs[job_id]['status'] = 'error'
          jobs[job_id]['error'] = str(e)
        return

    # Persist zip to temp file
    zip_buffer.seek(0)
    tmp_path = os.path.join(tempfile.gettempdir(), f"{job_id}.zip")
    with open(tmp_path, 'wb') as f:
      f.write(zip_buffer.getvalue())

    with jobs_lock:
      jobs[job_id]['status'] = 'done'
      jobs[job_id]['path'] = tmp_path
  except Exception as e:
    with jobs_lock:
      jobs[job_id]['status'] = 'error'
      jobs[job_id]['error'] = str(e)


@app.route('/start', methods=['POST'])
def start_job():
  """Start an async processing job. Returns a JSON job id and total files.
  Client should poll `/status/<job_id>` and then download from `/download/<job_id>`.
  """
  files = request.files.getlist('files')
  if not files:
    return jsonify({'error': 'No files uploaded'}), 400

  output_format = request.form.get('format', 'webp')
  try:
    quality = int(request.form.get('quality', 85))
  except Exception:
    quality = 85
  max_w = request.form.get('max_width') or None
  max_h = request.form.get('max_height') or None

  # Read all files into memory (safe for small uploads). For large uploads store on disk.
  file_blobs = []
  items = []
  for f in files:
    try:
      blob = f.read()
      file_blobs.append((f.filename or 'image', blob))
      items.append({'name': f.filename or 'image', 'size': len(blob), 'status': 'queued'})
    except Exception:
      continue

  job_id = uuid.uuid4().hex
  with jobs_lock:
    jobs[job_id] = {'total': len(file_blobs), 'processed': 0, 'status': 'processing', 'error': None, 'path': None, 'items': items}

  # Start background worker
  combine_pdf = request.form.get('combine_pdf', '0') in ('1', 'true', 'True')
  ocr = request.form.get('ocr', '0') in ('1', 'true', 'True')
  preserve_layout = request.form.get('preserve_layout', '0') in ('1', 'true', 'True')
  t = threading.Thread(target=_process_job, args=(job_id, file_blobs, output_format, quality, max_w, max_h, combine_pdf, ocr, preserve_layout), daemon=True)
  t.start()

  return jsonify({'job_id': job_id, 'total': len(file_blobs)})


@app.route('/status/<job_id>', methods=['GET'])
def job_status(job_id):
  with jobs_lock:
    job = jobs.get(job_id)
    if not job:
      return jsonify({'error': 'Job not found'}), 404
    return jsonify({'total': job['total'], 'processed': job['processed'], 'status': job['status'], 'error': job.get('error'), 'items': job.get('items', [])})


@app.route('/download/<job_id>', methods=['GET'])
def download_job(job_id):
  with jobs_lock:
    job = jobs.get(job_id)
    if not job:
      return Response('Job not found', status=404)
    if job['status'] != 'done' or not job.get('path'):
      return Response('Job not ready', status=400)
    path = job['path']

  # Send the ZIP file
  try:
    return send_file(path, mimetype='application/zip', as_attachment=True, download_name='optimized_images.zip')
  finally:
    # Cleanup after sending
    try:
      os.remove(path)
    except Exception:
      pass
    with jobs_lock:
      jobs.pop(job_id, None)


if __name__ == '__main__':
    # Run development server on port 8000
    # For production use a WSGI server like gunicorn or waitress
    app.run(host='127.0.0.1', port=8000, debug=True)